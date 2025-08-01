# src/core/document_processor.py
# HackRx 6.0 - Document Processing Module for PDF/DOCX parsing

import asyncio
import io
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import httpx

from ..utils.config import get_settings
from ..utils.logger import get_logger, log_performance_metric
from ..utils.helpers import download_file, clean_text, get_file_extension, generate_hash

logger = get_logger(__name__)


@dataclass
class DocumentChunk:
    """
    Represents a chunk of processed document text
    """
    text: str
    page_number: Optional[int] = None
    chunk_index: int = 0
    start_char: int = 0
    end_char: int = 0
    metadata: Dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass  
class ProcessedDocument:
    """
    Represents a fully processed document with chunks and metadata
    """
    url: str
    title: str
    total_pages: int
    total_characters: int
    chunks: List[DocumentChunk]
    document_hash: str
    file_type: str
    processing_time: float
    metadata: Dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentProcessor:
    """
    Main document processing class for handling PDF and DOCX files
    """
    
    def __init__(self):
        self.settings = get_settings()
        
    async def process_document_url(self, url: str) -> ProcessedDocument:
        """
        Download and process a document from URL
        
        Args:
            url: URL of the document to process
        
        Returns:
            ProcessedDocument with extracted text and chunks
        
        Raises:
            ValueError: If unsupported file type or processing fails
        """
        start_time = asyncio.get_event_loop().time()
        
        try:
            logger.info(f"Processing document from URL: {url}")
            
            # Validate file type
            file_extension = get_file_extension(url)
            if file_extension not in ['pdf', 'docx', 'doc']:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Download file
            file_content = await download_file(url, self.settings.max_document_size_mb)
            logger.info(f"Downloaded {len(file_content)} bytes")
            
            # Process based on file type
            if file_extension == 'pdf':
                processed_doc = await self._process_pdf(file_content, url)
            elif file_extension in ['docx', 'doc']:
                processed_doc = await self._process_docx(file_content, url)
            else:
                raise ValueError(f"Processing not implemented for {file_extension}")
            
            # Calculate processing time
            processing_time = asyncio.get_event_loop().time() - start_time
            processed_doc.processing_time = processing_time
            
            log_performance_metric(
                "document_processing", 
                processing_time, 
                len(processed_doc.chunks)
            )
            
            logger.info(
                f"Document processed successfully: "
                f"{len(processed_doc.chunks)} chunks, "
                f"{processed_doc.total_characters} characters, "
                f"{processing_time:.2f}s"
            )
            
            return processed_doc
            
        except Exception as e:
            processing_time = asyncio.get_event_loop().time() - start_time
            logger.error(f"Document processing failed after {processing_time:.2f}s: {e}")
            raise
    
    async def _process_pdf(self, file_content: bytes, url: str) -> ProcessedDocument:
        """
        Process PDF content and extract text
        
        Args:
            file_content: PDF file content as bytes
            url: Original URL for metadata
        
        Returns:
            ProcessedDocument with extracted content
        """
        try:
            # Open PDF from bytes
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            
            # Extract metadata
            metadata = pdf_document.metadata
            title = metadata.get('title', '') or Path(url).name
            total_pages = pdf_document.page_count
            
            logger.info(f"PDF opened: {total_pages} pages, title: '{title}'")
            
            # Extract text from all pages
            full_text = ""
            page_texts = []
            
            for page_num in range(total_pages):
                page = pdf_document[page_num]
                page_text = page.get_text()
                
                if page_text.strip():  # Only add non-empty pages
                    cleaned_text = clean_text(page_text)
                    page_texts.append((page_num + 1, cleaned_text))
                    full_text += cleaned_text + "\n\n"
            
            pdf_document.close()
            
            # Create document hash
            document_hash = generate_hash(full_text)
            
            # Split into chunks
            chunks = await self._create_chunks(full_text, page_texts)
            
            return ProcessedDocument(
                url=url,
                title=title,
                total_pages=total_pages,
                total_characters=len(full_text),
                chunks=chunks,
                document_hash=document_hash,
                file_type="pdf",
                processing_time=0.0,  # Will be set by caller
                metadata={
                    "pdf_metadata": metadata,
                    "pages_with_text": len(page_texts)
                }
            )
            
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise ValueError(f"Failed to process PDF: {e}")
    
    async def _process_docx(self, file_content: bytes, url: str) -> ProcessedDocument:
        """
        Process DOCX content and extract text
        
        Args:
            file_content: DOCX file content as bytes
            url: Original URL for metadata
        
        Returns:
            ProcessedDocument with extracted content
        """
        try:
            # Open DOCX from bytes
            docx_document = DocxDocument(io.BytesIO(file_content))
            
            # Extract text from paragraphs
            paragraphs = []
            full_text = ""
            
            for paragraph in docx_document.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:  # Only add non-empty paragraphs
                    cleaned_text = clean_text(para_text)
                    paragraphs.append(cleaned_text)
                    full_text += cleaned_text + "\n\n"
            
            # Extract title (try core properties first)
            title = ""
            try:
                if docx_document.core_properties.title:
                    title = docx_document.core_properties.title
                else:
                    title = Path(url).name
            except:
                title = Path(url).name
            
            logger.info(f"DOCX processed: {len(paragraphs)} paragraphs, title: '{title}'")
            
            # Create document hash
            document_hash = generate_hash(full_text)
            
            # Split into chunks (treating each paragraph as a "page")
            page_texts = [(i+1, para) for i, para in enumerate(paragraphs)]
            chunks = await self._create_chunks(full_text, page_texts)
            
            return ProcessedDocument(
                url=url,
                title=title,
                total_pages=len(paragraphs),  # Treat paragraphs as "pages"
                total_characters=len(full_text),
                chunks=chunks,
                document_hash=document_hash,
                file_type="docx",
                processing_time=0.0,  # Will be set by caller
                metadata={
                    "paragraph_count": len(paragraphs),
                    "has_title": bool(title and title != Path(url).name)
                }
            )
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise ValueError(f"Failed to process DOCX: {e}")
    
    async def _create_chunks(self, full_text: str, page_texts: List[Tuple[int, str]]) -> List[DocumentChunk]:
        """
        Split document text into chunks for processing
        
        Args:
            full_text: Complete document text
            page_texts: List of (page_number, page_text) tuples
        
        Returns:
            List of DocumentChunk objects
        """
        chunks = []
        chunk_size = self.settings.chunk_size
        overlap = self.settings.chunk_overlap
        
        # Strategy 1: Create chunks by sliding window over full text
        if len(full_text) <= chunk_size:
            # Small document - single chunk
            chunks.append(DocumentChunk(
                text=full_text,
                page_number=None,
                chunk_index=0,
                start_char=0,
                end_char=len(full_text),
                metadata={"chunk_type": "full_document"}
            ))
        else:
            # Large document - sliding window chunks
            start = 0
            chunk_index = 0
            
            while start < len(full_text):
                end = min(start + chunk_size, len(full_text))
                
                # Try to end at sentence boundary
                if end < len(full_text):
                    # Look for sentence endings near the target end
                    sentence_ends = ['.', '!', '?', '\n']
                    for i in range(end, max(start + chunk_size//2, 0), -1):
                        if full_text[i-1] in sentence_ends:
                            end = i
                            break
                
                chunk_text = full_text[start:end].strip()
                
                if chunk_text:  # Only add non-empty chunks
                    # Find which page(s) this chunk spans
                    page_numbers = self._find_pages_for_chunk(start, end, page_texts, full_text)
                    
                    chunks.append(DocumentChunk(
                        text=chunk_text,
                        page_number=page_numbers[0] if page_numbers else None,
                        chunk_index=chunk_index,
                        start_char=start,
                        end_char=end,
                        metadata={
                            "chunk_type": "sliding_window",
                            "pages": page_numbers,
                            "character_count": len(chunk_text)
                        }
                    ))
                    
                    chunk_index += 1
                
                # Move start position with overlap
                start = max(start + chunk_size - overlap, end)
                
                # Prevent infinite loop
                if start >= len(full_text):
                    break
        
        logger.info(f"Created {len(chunks)} chunks from document")
        return chunks
    
    def _find_pages_for_chunk(self, start_char: int, end_char: int, 
                             page_texts: List[Tuple[int, str]], full_text: str) -> List[int]:
        """
        Find which pages a chunk spans
        
        Args:
            start_char: Chunk start character position
            end_char: Chunk end character position  
            page_texts: List of (page_number, page_text) tuples
            full_text: Complete document text
        
        Returns:
            List of page numbers that overlap with the chunk
        """
        pages = []
        current_pos = 0
        
        for page_num, page_text in page_texts:
            page_start = current_pos
            page_end = current_pos + len(page_text) + 2  # +2 for "\n\n" separator
            
            # Check if chunk overlaps with this page
            if not (end_char <= page_start or start_char >= page_end):
                pages.append(page_num)
            
            current_pos = page_end
        
        return pages
    
    async def extract_key_sections(self, processed_doc: ProcessedDocument) -> Dict[str, List[str]]:
        """
        Extract key sections from document (insurance/legal specific)
        
        Args:
            processed_doc: Processed document to analyze
        
        Returns:
            Dictionary mapping section types to content lists
        """
        sections = {
            "definitions": [],
            "coverage": [],
            "exclusions": [],
            "conditions": [],
            "claims": [],
            "waiting_periods": [],
            "benefits": []
        }
        
        # Keywords for section identification
        section_keywords = {
            "definitions": ["definition", "means", "refers to", "shall mean"],
            "coverage": ["coverage", "covered", "benefit", "indemnify", "shall cover"],
            "exclusions": ["exclusion", "excluded", "not covered", "shall not"],
            "conditions": ["condition", "terms", "requirement", "provided that"],
            "claims": ["claim", "procedure", "settlement", "reimbursement"],
            "waiting_periods": ["waiting period", "waiting", "months", "days"],
            "benefits": ["benefit", "limit", "maximum", "sum insured"]
        }
        
        # Analyze each chunk
        for chunk in processed_doc.chunks:
            text_lower = chunk.text.lower()
            
            for section_type, keywords in section_keywords.items():
                # Check if chunk contains relevant keywords
                keyword_count = sum(1 for keyword in keywords if keyword in text_lower)
                
                if keyword_count >= 2:  # Require multiple keyword matches
                    sections[section_type].append(chunk.text)
        
        # Log section analysis
        section_counts = {k: len(v) for k, v in sections.items() if v}
        if section_counts:
            logger.info(f"Extracted sections: {section_counts}")
        
        return sections